import os
import json
import re
import logging
from pathlib import Path
from typing import Optional
from dataclasses import dataclass, field
from langchain_openai import ChatOpenAI
from langchain_groq import ChatGroq


from dotenv import load_dotenv
from langchain_google_genai import ChatGoogleGenerativeAI

try:
    import fitz  # PyMuPDF
except ImportError:
    import pymupdf as fitz

import pdfplumber
import docx

load_dotenv()
logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────
# Constants
# ─────────────────────────────────────────────

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}
MAX_TEXT_LENGTH = 10_000
COLUMN_SPLIT_RATIO = 0.5  # fraction of page width to split left/right columns

RESUME_JSON_SCHEMA = {
    "name": "",
    "email": "",
    "phone": "",
    "location": "",
    "linkedin": "",
    "github": "",
    "skills": [],
    "experience": [],
    "education": [],
    "certifications": [],
    "projects": [],
    "publications": [],
    "summary": "",
    "years_of_experience": 0,
}

PARSE_PROMPT_TEMPLATE = """
Parse this resume and return ONLY valid JSON. Handle ambiguous formatting intelligently.

Resume text:
{resume_text}

Required JSON structure:
{{
  "name": "Full name or empty string",
  "email": "email@example.com or empty",
  "phone": "+1234567890 or empty",
  "location": "City, Country or empty",
  "linkedin": "linkedin url or empty",
  "github": "github url or empty",
  "skills": ["skill1", "skill2"],
  "experience": [
    {{"company": "Company Name", "role": "Job Title", "duration": "MM/YYYY - MM/YYYY", "responsibilities": ["responsibility1"]}}
  ],
  "education": [
    {{"institution": "University", "degree": "B.Sc.", "field": "Computer Science", "year": "2020"}}
  ],
  "certifications": ["Certification Name"],
  "projects": [
    {{"name": "Project Name", "description": "Brief description", "technologies": ["tech1"]}}
  ],
  "publications": ["Publication title"],
  "summary": "Brief professional summary",
  "years_of_experience": 0
}}

Rules:
- Extract ALL information you can find
- For missing fields, use empty strings/arrays
- Infer years of experience from work history
- Return ONLY valid JSON, no markdown
- Do not include any explanatory text outside the JSON
"""

# ─────────────────────────────────────────────
# Extraction result
# ─────────────────────────────────────────────

@dataclass
class ExtractionResult:
    text: str
    layout_type: str
    is_multi_column: bool = False
    pages: int = 0
    warnings: list[str] = field(default_factory=list)


# ─────────────────────────────────────────────
# Extractors
# ─────────────────────────────────────────────

def _extract_pdf(file_path: Path) -> ExtractionResult:
    """Extract text from PDF, handling multi-column layouts via PyMuPDF."""
    pages_text = []
    is_multi_column = False

    try:
        doc = fitz.open(str(file_path))
        for page in doc:
            left_spans, right_spans = [], []
            mid_x = page.rect.width * COLUMN_SPLIT_RATIO

            try:
                blocks = page.get_text("dict").get("blocks", [])
                for block in blocks:
                    for line in block.get("lines", []):
                        for span in line.get("spans", []):
                            text = span.get("text", "").strip()
                            if not text:
                                continue
                            (left_spans if span["bbox"][0] < mid_x else right_spans).append(text)

                if left_spans and right_spans:
                    is_multi_column = True
                    # Read left column first, then right (natural reading order)
                    pages_text.append(" ".join(left_spans) + "\n" + " ".join(right_spans))
                else:
                    pages_text.append(page.get_text())
            except Exception:
                pages_text.append(page.get_text())

        doc.close()
        return ExtractionResult(
            text="\n".join(pages_text),
            layout_type="pdf-pymupdf",
            is_multi_column=is_multi_column,
            pages=len(pages_text),
        )

    except Exception as e:
        logger.warning(f"PyMuPDF failed ({e}), falling back to pdfplumber")

    # Fallback: pdfplumber
    pages_text = []
    with pdfplumber.open(str(file_path)) as pdf:
        for page in pdf.pages:
            pages_text.append(page.extract_text() or "")

    return ExtractionResult(
        text="\n".join(pages_text),
        layout_type="pdf-pdfplumber",
        pages=len(pages_text),
        warnings=["Used pdfplumber fallback; multi-column detection unavailable"],
    )


def _extract_docx(file_path: Path) -> ExtractionResult:
    """
    Extract text from DOCX, including multi-column sections via XML.
    DOCX columns are defined per section; we read all runs in order.
    """
    doc = docx.Document(str(file_path))
    is_multi_column = False

    # Check for multi-column sections via XML namespace
    try:
        from lxml import etree
        W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        for section in doc.sections:
            cols_el = section._sectPr.find(f"{{{W_NS}}}cols")
            if cols_el is not None:
                num_cols = int(cols_el.get(f"{{{W_NS}}}num", 1))
                if num_cols > 1:
                    is_multi_column = True
                    break
    except Exception:
        pass  # lxml unavailable or XML parse error — skip detection

    # Extract all paragraph text in document order
    # (Word stores multi-column text linearly in XML, so order is preserved)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return ExtractionResult(
        text="\n".join(paragraphs),
        layout_type="docx",
        is_multi_column=is_multi_column,
    )


def _extract_txt(file_path: Path) -> ExtractionResult:
    """
    Extract text from plain .txt, with heuristic multi-column detection.
    Detects columns by checking if lines have a large mid-line gap (multiple spaces).
    """
    text = file_path.read_text(encoding="utf-8", errors="ignore")
    lines = text.splitlines()
    is_multi_column = False
    reconstructed = []

    for line in lines:
        # Heuristic: 3+ consecutive spaces mid-line = likely two columns
        match = re.search(r"(\S.*?)\s{3,}(\S.*)", line)
        if match:
            is_multi_column = True
            reconstructed.append(match.group(1).strip())
            reconstructed.append(match.group(2).strip())
        else:
            reconstructed.append(line)

    return ExtractionResult(
        text="\n".join(reconstructed),
        layout_type="txt",
        is_multi_column=is_multi_column,
    )


# ─────────────────────────────────────────────
# Main extractor dispatcher
# ─────────────────────────────────────────────

def extract_text(file_path: Path) -> ExtractionResult:
    ext = file_path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type '{ext}'. Supported: {SUPPORTED_EXTENSIONS}")

    extractors = {
        ".pdf": _extract_pdf,
        ".docx": _extract_docx,
        ".txt": _extract_txt,
    }
    result = extractors[ext](file_path)

    if not result.text.strip():
        raise ValueError(f"No text could be extracted from '{file_path.name}'")

    return result


# ─────────────────────────────────────────────
# LLM helpers
# ─────────────────────────────────────────────

def _clean_llm_output(raw: str) -> str:
    """Strip markdown fences and extraneous text; return bare JSON string."""
    text = raw.strip()

    # Remove ```json ... ``` or ``` ... ```
    text = re.sub(r"^```[a-zA-Z]*\n?", "", text)
    text = re.sub(r"\n?```$", "", text)
    text = text.strip()

    # If still not starting with {, try to extract JSON object
    if not text.startswith("{"):
        match = re.search(r"\{.*\}", text, re.DOTALL)
        if match:
            text = match.group()

    return text


def _ensure_schema(data: dict) -> dict:
    """Ensure all required fields exist with correct default types."""
    for key, default in RESUME_JSON_SCHEMA.items():
        if key not in data:
            data[key] = default if not isinstance(default, list) else []
    return data


# ─────────────────────────────────────────────
# ResumeParser
# ─────────────────────────────────────────────

class ResumeParser:
    def __init__(self, model: str = "llama-3.3-70b-versatile"):
        self.llm = ChatGroq(
            model=model,
            api_key=os.getenv("GROQ_API_KEY"),
            temperature=0,
        )

    def parse_resume(self, file_path: str | Path) -> dict:
        file_path = Path(file_path)

        if not file_path.exists():
            return self._error_result(f"File not found: {file_path}")

        # 1. Extract text
        try:
            extraction = extract_text(file_path)
        except Exception as e:
            logger.error(f"Extraction failed: {e}")
            return self._error_result(str(e))

        logger.info(
            f"Extracted {len(extraction.text)} chars | layout={extraction.layout_type} "
            f"| multi_column={extraction.is_multi_column}"
        )

        # 2. Truncate if needed
        raw_text = extraction.text
        truncated = False
        if len(raw_text) > MAX_TEXT_LENGTH:
            raw_text = raw_text[:MAX_TEXT_LENGTH] + "\n... (truncated)"
            truncated = True
            logger.warning("Resume text truncated to fit token limit")

        # 3. Call LLM
        prompt = PARSE_PROMPT_TEMPLATE.format(resume_text=raw_text)
        try:
            response = self.llm.invoke(prompt)
            cleaned = _clean_llm_output(response.content)
            parsed = json.loads(cleaned)
        except json.JSONDecodeError as e:
            logger.error(f"JSON decode error: {e}")
            raw_out = response.content if "response" in dir() else "No LLM output"
            return self._error_result(f"JSON parsing failed: {e}", raw_output=raw_out)
        except Exception as e:
            logger.error(f"LLM call failed: {e}")
            return self._error_result(str(e))

        # 4. Normalize and attach metadata
        parsed = _ensure_schema(parsed)
        parsed["_metadata"] = {
            "source_file": file_path.name,
            "layout_type": extraction.layout_type,
            "is_multi_column": extraction.is_multi_column,
            "pages": extraction.pages,
            "raw_text_length": len(extraction.text),
            "truncated": truncated,
            "warnings": extraction.warnings,
        }

        return parsed

    @staticmethod
    def _error_result(message: str, raw_output: str = "") -> dict:
        return {
            "error": message,
            "raw_output": raw_output,
            **{k: ([] if isinstance(v, list) else "") for k, v in RESUME_JSON_SCHEMA.items()},
        }


# ─────────────────────────────────────────────
# Module-level convenience function
# ─────────────────────────────────────────────

_parser_instance: Optional[ResumeParser] = None

def parse_resume(file_path: str) -> dict:
    global _parser_instance
    if _parser_instance is None:
        _parser_instance = ResumeParser()
    return _parser_instance.parse_resume(file_path)
